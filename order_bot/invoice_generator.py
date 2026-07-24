from __future__ import annotations

import csv
import html
import os
import re
import tempfile
import zipfile
from base64 import b64encode
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree

from .paths import browser_cache_dir


HEADER_MAP = {
    "vat invoice number": "vat_invoice_number",
    "payment date": "payment_date",
    "date of supply": "date_of_supply",
    "supplier": "supplier",
    "sku": "sku",
    "description": "description",
    "specification": "specification",
    "quantity": "quantity",
    "unit price": "unit_price",
    "net amount": "net_amount",
    "vat amount": "vat_amount",
    "gross amount": "gross_amount",
    "company name": "company_name",
    "company": "company_name",
    "公司名称": "company_name",
}

INVOICE_TEMPLATE_SUFFIXES = (".pdf", ".html", ".htm")
INVOICE_OUTPUT_FORMAT_PDF = "pdf"
INVOICE_OUTPUT_FORMAT_HTML = "html"
INVOICE_OUTPUT_FORMAT_PNG = "png"
INVOICE_OUTPUT_FORMAT_JPG = "jpg"
INVOICE_OUTPUT_FORMAT_DOCX = "docx"
INVOICE_OUTPUT_FORMATS = (
    INVOICE_OUTPUT_FORMAT_PDF,
    INVOICE_OUTPUT_FORMAT_HTML,
    INVOICE_OUTPUT_FORMAT_PNG,
    INVOICE_OUTPUT_FORMAT_JPG,
    INVOICE_OUTPUT_FORMAT_DOCX,
)
INVOICE_OUTPUT_FORMAT_LABELS = {
    INVOICE_OUTPUT_FORMAT_PDF: "PDF (.pdf)",
    INVOICE_OUTPUT_FORMAT_HTML: "HTML (.html)",
    INVOICE_OUTPUT_FORMAT_PNG: "图片 PNG (.png)",
    INVOICE_OUTPUT_FORMAT_JPG: "图片 JPG (.jpg)",
    INVOICE_OUTPUT_FORMAT_DOCX: "Word (.docx)",
}
INVOICE_PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

COMPANY_NAME_COLUMN = "company_name"
KNOWN_COLUMNS = {
    "company_name",
    "vat_invoice_number",
    "payment_date",
    "date_of_supply",
    "supplier",
    "sku",
    "description",
    "specification",
    "quantity",
    "unit_price",
    "net_amount",
    "vat_amount",
    "gross_amount",
}


@dataclass(frozen=True)
class InvoiceDataRow:
    row_index: int
    company_name: str
    values: dict[str, str]


@dataclass(frozen=True)
class InvoiceTemplateAnalysis:
    path: Path
    style: str
    page_count: int
    summary: str
    field_names: tuple[str, ...] = ()


@dataclass(frozen=True)
class InvoiceGenerationResult:
    row_index: int
    company_name: str
    invoice_number: str
    output_path: Path


def normalize_invoice_header(value: str) -> str:
    raw = (value or "").strip()
    if not raw:
        return ""
    mapped = HEADER_MAP.get(raw.casefold())
    if mapped:
        return mapped
    normalized = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "_", raw).strip("_").lower()
    normalized = re.sub(r"_+", "_", normalized)
    return HEADER_MAP.get(normalized.replace("_", " ").casefold(), normalized)


def invoice_output_format_labels() -> tuple[str, ...]:
    return tuple(INVOICE_OUTPUT_FORMAT_LABELS[output_format] for output_format in INVOICE_OUTPUT_FORMATS)


def invoice_output_format_from_label(value: str) -> str:
    normalized = (value or "").strip().casefold()
    for output_format, label in INVOICE_OUTPUT_FORMAT_LABELS.items():
        if normalized in {output_format, label.casefold()}:
            return output_format
    raise ValueError(f"不支持的发票导出格式：{value}")


def invoice_output_extension(output_format: str) -> str:
    output_format = invoice_output_format_from_label(output_format)
    return ".jpg" if output_format == INVOICE_OUTPUT_FORMAT_JPG else f".{output_format}"


def load_invoice_rows(path: Path) -> list[InvoiceDataRow]:
    rows = _load_raw_rows(path)
    header_index = _find_header_index(rows)
    if header_index is None:
        raise ValueError("数据文件未找到发票表头，请确认包含 company_name 或 supplier 列。")

    headers = [normalize_invoice_header(value) for value in rows[header_index]]
    data_rows: list[InvoiceDataRow] = []
    for offset, raw_row in enumerate(rows[header_index + 1 :], start=header_index + 2):
        values: dict[str, str] = {}
        for column_index, header in enumerate(headers):
            if not header:
                continue
            values[header] = raw_row[column_index].strip() if column_index < len(raw_row) else ""
        if not any(values.values()):
            continue
        if not values.get(COMPANY_NAME_COLUMN) and values.get("supplier"):
            values[COMPANY_NAME_COLUMN] = values["supplier"]
        company_name = values.get(COMPANY_NAME_COLUMN, "").strip()
        if not company_name:
            raise ValueError(f"第 {offset} 行缺少 company_name，公司名称必须作为模板匹配 key。")
        data_rows.append(InvoiceDataRow(row_index=offset, company_name=company_name, values=values))

    if not data_rows:
        raise ValueError("数据文件没有可生成发票的数据行。")
    return data_rows


def unique_companies(rows: list[InvoiceDataRow]) -> list[str]:
    seen: set[str] = set()
    companies: list[str] = []
    for row in rows:
        key = normalize_company_key(row.company_name)
        if key in seen:
            continue
        seen.add(key)
        companies.append(row.company_name)
    return companies


def normalize_company_key(value: str) -> str:
    return re.sub(r"[^0-9a-z]+", "", (value or "").casefold())


def analyze_invoice_template(path: Path) -> InvoiceTemplateAnalysis:
    if path.suffix.lower() in {".html", ".htm"}:
        template_text = read_invoice_html_template(path)
        field_names = tuple(extract_invoice_placeholders(template_text))
        summary = re.sub(r"<[^>]+>", " ", template_text)
        summary = " ".join(html.unescape(summary).split())[:200]
        return InvoiceTemplateAnalysis(
            path=path,
            style="html",
            page_count=1,
            summary=summary,
            field_names=field_names,
        )

    reader = _pdf_reader(path)
    text = "\n".join(page.extract_text() or "" for page in reader.pages[:2])
    style = _detect_template_style(path, text)
    field_names = tuple(_pdf_form_field_names(reader))
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    summary = " / ".join(lines[:4])
    return InvoiceTemplateAnalysis(
        path=path,
        style=style,
        page_count=len(reader.pages),
        summary=summary,
        field_names=field_names,
    )


def generate_invoice_pdf(
    *,
    template_path: Path,
    row: InvoiceDataRow,
    output_path: Path,
) -> InvoiceGenerationResult:
    reader = _pdf_reader(template_path)
    writer = _pdf_writer()
    if not reader.pages:
        raise ValueError(f"发票模板没有页面：{template_path}")

    if _pdf_has_form_fields(reader):
        _generate_pdf_from_form_fields(reader=reader, row=row, output_path=output_path)
        return InvoiceGenerationResult(
            row_index=row.row_index,
            company_name=row.company_name,
            invoice_number=_value(row.values, "vat_invoice_number", "invoice_number"),
            output_path=output_path,
        )

    first_page = reader.pages[0]
    width = float(first_page.mediabox.width)
    height = float(first_page.mediabox.height)
    text = first_page.extract_text() or ""
    style = _detect_template_style(template_path, text)
    overlay_page = _build_overlay_page(style, row.values, width, height)
    first_page.merge_page(overlay_page)
    writer.add_page(first_page)
    for page in reader.pages[1:]:
        writer.add_page(page)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("wb") as handle:
        writer.write(handle)

    return InvoiceGenerationResult(
        row_index=row.row_index,
        company_name=row.company_name,
        invoice_number=_value(row.values, "vat_invoice_number", "invoice_number"),
        output_path=output_path,
    )


def generate_invoice_file(
    *,
    template_path: Path,
    row: InvoiceDataRow,
    output_path: Path,
    output_format: str,
) -> InvoiceGenerationResult:
    output_format = invoice_output_format_from_label(output_format)
    template_suffix = template_path.suffix.lower()
    if template_suffix not in INVOICE_TEMPLATE_SUFFIXES:
        raise ValueError("发票模板仅支持 .pdf、.html、.htm。")

    if template_suffix in {".html", ".htm"}:
        return _generate_invoice_from_html_template(
            template_path=template_path,
            row=row,
            output_path=output_path,
            output_format=output_format,
        )

    if output_format == INVOICE_OUTPUT_FORMAT_PDF:
        return generate_invoice_pdf(template_path=template_path, row=row, output_path=output_path)

    with tempfile.TemporaryDirectory(prefix="auto-order-invoice-") as temp_dir:
        temp_pdf = Path(temp_dir) / "invoice.pdf"
        generate_invoice_pdf(template_path=template_path, row=row, output_path=temp_pdf)
        _export_pdf_to_format(temp_pdf, output_path, output_format)

    return InvoiceGenerationResult(
        row_index=row.row_index,
        company_name=row.company_name,
        invoice_number=_value(row.values, "vat_invoice_number", "invoice_number"),
        output_path=output_path,
    )


def output_invoice_path(output_dir: Path, row: InvoiceDataRow, *, output_format: str = INVOICE_OUTPUT_FORMAT_PDF) -> Path:
    invoice_number = _value(row.values, "vat_invoice_number", "invoice_number") or f"row-{row.row_index}"
    company_name = safe_filename(row.company_name)
    filename = f"{company_name}-{safe_filename(invoice_number)}{invoice_output_extension(output_format)}"
    return output_dir / company_name / filename


def safe_filename(value: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "-", (value or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    return cleaned[:120] or datetime.now().strftime("invoice-%Y%m%d-%H%M%S")


def extract_invoice_placeholders(template_text: str) -> list[str]:
    seen: set[str] = set()
    placeholders: list[str] = []
    for match in INVOICE_PLACEHOLDER_PATTERN.finditer(template_text or ""):
        name = match.group(1).strip()
        if name and name not in seen:
            seen.add(name)
            placeholders.append(name)
    return placeholders


def read_invoice_html_template(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def render_invoice_html_template(template_text: str, row: dict[str, str]) -> str:
    def replace_match(match: re.Match[str]) -> str:
        placeholder = match.group(1).strip()
        value = _value_for_field_name(row, placeholder)
        return value if value else match.group(0)

    return INVOICE_PLACEHOLDER_PATTERN.sub(replace_match, template_text)


def _load_raw_rows(path: Path) -> list[list[str]]:
    suffix = path.suffix.lower()
    if suffix == ".csv":
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            return [[(value or "").strip() for value in row] for row in csv.reader(handle)]
    if suffix == ".xlsx":
        return _load_xlsx_rows(path)
    raise ValueError("发票数据文件仅支持 .csv 或 .xlsx。")


def _find_header_index(rows: list[list[str]]) -> int | None:
    best_index: int | None = None
    best_score = 0
    for index, row in enumerate(rows[:20]):
        normalized = {normalize_invoice_header(value) for value in row if value}
        score = len(normalized & KNOWN_COLUMNS)
        if score > best_score:
            best_score = score
            best_index = index
    if best_score >= 2:
        return best_index
    return None


def _load_xlsx_rows(path: Path) -> list[list[str]]:
    ns = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        shared_strings = _read_shared_strings(archive, ns)
        sheet_name = _first_sheet_path(archive)
        root = ElementTree.fromstring(archive.read(sheet_name))

    rows: list[list[str]] = []
    for row_element in root.findall(".//main:sheetData/main:row", ns):
        values: list[str] = []
        current_col = 0
        for cell in row_element.findall("main:c", ns):
            cell_ref = cell.attrib.get("r", "")
            target_col = _column_index_from_cell_ref(cell_ref)
            while current_col < target_col:
                values.append("")
                current_col += 1
            values.append(_cell_text(cell, shared_strings, ns))
            current_col += 1
        rows.append(values)
    return rows


def _read_shared_strings(archive: zipfile.ZipFile, ns: dict[str, str]) -> list[str]:
    if "xl/sharedStrings.xml" not in archive.namelist():
        return []
    root = ElementTree.fromstring(archive.read("xl/sharedStrings.xml"))
    values = []
    for item in root.findall("main:si", ns):
        text_parts = [node.text or "" for node in item.findall(".//main:t", ns)]
        values.append("".join(text_parts))
    return values


def _first_sheet_path(archive: zipfile.ZipFile) -> str:
    workbook = ElementTree.fromstring(archive.read("xl/workbook.xml"))
    rels_root = ElementTree.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
    ns = {
        "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
        "rel": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pkg": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    first_sheet = workbook.find("main:sheets/main:sheet", ns)
    if first_sheet is None:
        raise ValueError("xlsx 没有工作表。")
    rel_id = first_sheet.attrib[f"{{{ns['rel']}}}id"]
    for relationship in rels_root.findall("pkg:Relationship", ns):
        if relationship.attrib.get("Id") == rel_id:
            target = relationship.attrib["Target"].lstrip("/")
            return target if target.startswith("xl/") else f"xl/{target}"
    raise ValueError("xlsx 工作表关系缺失。")


def _cell_text(cell, shared_strings: list[str], ns: dict[str, str]) -> str:
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        return "".join(node.text or "" for node in cell.findall(".//main:t", ns)).strip()
    value_node = cell.find("main:v", ns)
    if value_node is None or value_node.text is None:
        return ""
    raw_value = value_node.text
    if cell_type == "s":
        index = int(raw_value)
        return shared_strings[index].strip() if index < len(shared_strings) else ""
    return raw_value.strip()


def _column_index_from_cell_ref(cell_ref: str) -> int:
    letters = "".join(ch for ch in cell_ref if ch.isalpha())
    index = 0
    for letter in letters:
        index = index * 26 + ord(letter.upper()) - ord("A") + 1
    return max(0, index - 1)


def _pdf_reader(path: Path):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("缺少 PDF 处理依赖 pypdf，请重新安装依赖或使用最新版安装包。") from exc
    return PdfReader(str(path))


def _pdf_writer():
    try:
        from pypdf import PdfWriter
    except ImportError as exc:
        raise RuntimeError("缺少 PDF 处理依赖 pypdf，请重新安装依赖或使用最新版安装包。") from exc
    return PdfWriter()


def _generate_invoice_from_html_template(
    *,
    template_path: Path,
    row: InvoiceDataRow,
    output_path: Path,
    output_format: str,
) -> InvoiceGenerationResult:
    template_text = read_invoice_html_template(template_path)
    rendered_html = render_invoice_html_template(template_text, row.values)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_format == INVOICE_OUTPUT_FORMAT_HTML:
        output_path.write_text(rendered_html, encoding="utf-8")
    elif output_format == INVOICE_OUTPUT_FORMAT_PDF:
        _render_html_to_pdf(rendered_html, output_path)
    elif output_format in {INVOICE_OUTPUT_FORMAT_PNG, INVOICE_OUTPUT_FORMAT_JPG}:
        _render_html_to_image(rendered_html, output_path, output_format)
    elif output_format == INVOICE_OUTPUT_FORMAT_DOCX:
        with tempfile.TemporaryDirectory(prefix="auto-order-html-docx-") as temp_dir:
            image_path = Path(temp_dir) / "invoice.png"
            _render_html_to_image(rendered_html, image_path, INVOICE_OUTPUT_FORMAT_PNG)
            _images_to_docx([image_path], output_path)
    else:
        raise ValueError(f"不支持的发票导出格式：{output_format}")

    return InvoiceGenerationResult(
        row_index=row.row_index,
        company_name=row.company_name,
        invoice_number=_value(row.values, "vat_invoice_number", "invoice_number"),
        output_path=output_path,
    )


def _export_pdf_to_format(pdf_path: Path, output_path: Path, output_format: str) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_format == INVOICE_OUTPUT_FORMAT_HTML:
        _pdf_to_html_images(pdf_path, output_path)
    elif output_format in {INVOICE_OUTPUT_FORMAT_PNG, INVOICE_OUTPUT_FORMAT_JPG}:
        _pdf_to_images(pdf_path, output_path, output_format)
    elif output_format == INVOICE_OUTPUT_FORMAT_DOCX:
        with tempfile.TemporaryDirectory(prefix="auto-order-pdf-docx-") as temp_dir:
            image_paths = _pdf_to_images(pdf_path, Path(temp_dir) / "invoice.png", INVOICE_OUTPUT_FORMAT_PNG)
            _images_to_docx(image_paths, output_path)
    else:
        raise ValueError(f"不支持从 PDF 导出为 {output_format}")


def _render_html_to_pdf(rendered_html: str, output_path: Path) -> None:
    def render(page) -> None:
        page.pdf(path=str(output_path), format="A4", print_background=True)

    _render_html_with_playwright(rendered_html, render)


def _render_html_to_image(rendered_html: str, output_path: Path, output_format: str) -> None:
    image_type = "jpeg" if output_format == INVOICE_OUTPUT_FORMAT_JPG else "png"

    def render(page) -> None:
        page.screenshot(path=str(output_path), full_page=True, type=image_type)

    _render_html_with_playwright(rendered_html, render)


def _render_html_with_playwright(rendered_html: str, render_callback) -> None:
    try:
        from playwright.sync_api import Error as PlaywrightError
        from playwright.sync_api import sync_playwright
    except ImportError as exc:
        raise RuntimeError("缺少 HTML 渲染依赖 Playwright，请重新安装依赖或使用最新版安装包。") from exc

    os.environ.setdefault("PLAYWRIGHT_BROWSERS_PATH", str(browser_cache_dir()))
    playwright = None
    browser = None
    try:
        playwright = sync_playwright().start()
        browser = playwright.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": 900, "height": 1200}, device_scale_factor=1)
        page.set_content(rendered_html, wait_until="networkidle", timeout=60000)
        render_callback(page)
    except PlaywrightError as exc:
        raise RuntimeError(
            "HTML 模板导出 PDF/图片/Word 需要 Chromium 浏览器。请先运行：python -m playwright install chromium"
        ) from exc
    finally:
        if browser is not None:
            browser.close()
        if playwright is not None:
            playwright.stop()


def _pdf_to_images(pdf_path: Path, output_path: Path, output_format: str) -> list[Path]:
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:
        raise RuntimeError("缺少 PDF 转图片依赖 pypdfium2，请重新安装依赖或使用最新版安装包。") from exc

    pdf = pdfium.PdfDocument(str(pdf_path))
    image_paths: list[Path] = []
    page_count = len(pdf)
    image_format = "JPEG" if output_format == INVOICE_OUTPUT_FORMAT_JPG else "PNG"
    suffix = ".jpg" if output_format == INVOICE_OUTPUT_FORMAT_JPG else ".png"
    for index in range(page_count):
        page = pdf[index]
        bitmap = page.render(scale=2).to_pil()
        page_output = output_path if page_count == 1 else output_path.with_name(f"{output_path.stem}-p{index + 1}{suffix}")
        page_output.parent.mkdir(parents=True, exist_ok=True)
        if image_format == "JPEG":
            bitmap = bitmap.convert("RGB")
            bitmap.save(page_output, format=image_format, quality=95)
        else:
            bitmap.save(page_output, format=image_format)
        image_paths.append(page_output)
    return image_paths


def _pdf_to_html_images(pdf_path: Path, output_path: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="auto-order-pdf-html-") as temp_dir:
        image_paths = _pdf_to_images(pdf_path, Path(temp_dir) / "invoice.png", INVOICE_OUTPUT_FORMAT_PNG)
        images = []
        for image_path in image_paths:
            encoded = b64encode(image_path.read_bytes()).decode("ascii")
            images.append(
                f'<img src="data:image/png;base64,{encoded}" '
                'style="display:block;width:100%;max-width:900px;margin:0 auto 18px;" />'
            )
        html_text = (
            "<!doctype html><html><head><meta charset=\"utf-8\">"
            "<title>Invoice</title></head><body style=\"margin:0;background:#f4f4f4;padding:24px;\">"
            + "\n".join(images)
            + "</body></html>"
        )
        output_path.write_text(html_text, encoding="utf-8")


def _images_to_docx(image_paths: list[Path], output_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError("缺少 Word 导出依赖 python-docx，请重新安装依赖或使用最新版安装包。") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    image_width = section.page_width - section.left_margin - section.right_margin
    for index, image_path in enumerate(image_paths):
        if index:
            document.add_page_break()
        document.add_picture(str(image_path), width=image_width)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _pdf_has_form_fields(reader) -> bool:
    return bool(_pdf_form_field_names(reader))


def _pdf_form_field_names(reader) -> list[str]:
    try:
        fields = reader.get_fields() or {}
    except Exception:
        return []
    return [str(name) for name in fields.keys() if str(name).strip()]


def _generate_pdf_from_form_fields(*, reader, row: InvoiceDataRow, output_path: Path) -> None:
    try:
        from pypdf.generic import ArrayObject, NameObject
    except ImportError as exc:
        raise RuntimeError("缺少 PDF 处理依赖 pypdf，请重新安装依赖或使用最新版安装包。") from exc

    writer = _pdf_writer()
    for source_page in reader.pages:
        writer.add_page(source_page)
        page = writer.pages[-1]
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        field_draws: list[tuple[float, float, float, float, str]] = []
        remaining_annotations = ArrayObject()

        annotations = page.get("/Annots") or []
        for annotation in annotations:
            annotation_object = annotation.get_object()
            if annotation_object.get("/Subtype") != "/Widget":
                remaining_annotations.append(annotation)
                continue

            field_name = _field_name_from_widget(annotation_object)
            value = _value_for_field_name(row.values, field_name)
            rect = annotation_object.get("/Rect")
            if value and rect and len(rect) >= 4:
                x1, y1, x2, y2 = [float(item) for item in rect[:4]]
                field_draws.append((x1, y1, x2 - x1, y2 - y1, value))

        if field_draws:
            overlay_page = _build_form_overlay_page(field_draws, width, height)
            page.merge_page(overlay_page)

        if remaining_annotations:
            page[NameObject("/Annots")] = remaining_annotations
        elif "/Annots" in page:
            del page["/Annots"]

    if "/AcroForm" in writer._root_object:
        del writer._root_object["/AcroForm"]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("wb") as handle:
        writer.write(handle)


def _field_name_from_widget(annotation_object) -> str:
    names: list[str] = []
    current = annotation_object
    seen: set[int] = set()
    while current and id(current) not in seen:
        seen.add(id(current))
        name = current.get("/T")
        if name:
            names.append(str(name))
        parent = current.get("/Parent")
        current = parent.get_object() if parent else None
    return ".".join(reversed(names))


def _value_for_field_name(row: dict[str, str], field_name: str) -> str:
    variable_name = _field_name_to_variable(field_name)
    if not variable_name:
        return ""
    value = _value(row, variable_name)
    if value:
        return value
    normalized_variable = normalize_invoice_header(variable_name)
    normalized_lookup = _normalize_lookup_key(normalized_variable)
    for key, row_value in row.items():
        if _normalize_lookup_key(key) == normalized_lookup and row_value:
            return row_value
    return ""


def _field_name_to_variable(field_name: str) -> str:
    name = (field_name or "").strip()
    if "." in name:
        name = name.rsplit(".", 1)[-1]
    placeholder_match = re.fullmatch(r"\{\{\s*([^{}]+?)\s*\}\}", name)
    if placeholder_match:
        name = placeholder_match.group(1).strip()
    return normalize_invoice_header(name)


def _normalize_lookup_key(value: str) -> str:
    return re.sub(r"[^0-9a-z\u4e00-\u9fff]+", "", (value or "").casefold())


def _build_form_overlay_page(fields: list[tuple[float, float, float, float, str]], width: float, height: float):
    try:
        from pypdf import PdfReader
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise RuntimeError("缺少 PDF 生成依赖 reportlab / pypdf，请重新安装依赖或使用最新版安装包。") from exc

    buffer = BytesIO()
    pdf_canvas = canvas.Canvas(buffer, pagesize=(width, height))
    for x, y, field_width, field_height, value in fields:
        font_size = max(7, min(13, int(field_height * 0.45)))
        text_x = x + 2
        text_y = y + max(2, (field_height - font_size) / 2)
        if field_width > 70:
            _wrap_text(pdf_canvas, value, text_x, text_y, max(field_width - 4, 20), size=font_size)
        else:
            _text(pdf_canvas, text_x, text_y, value, size=font_size)
    pdf_canvas.save()
    buffer.seek(0)
    return PdfReader(buffer).pages[0]


def _build_overlay_page(style: str, row: dict[str, str], width: float, height: float):
    try:
        from pypdf import PdfReader
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise RuntimeError("缺少 PDF 生成依赖 reportlab / pypdf，请重新安装依赖或使用最新版安装包。") from exc

    buffer = BytesIO()
    pdf_canvas = canvas.Canvas(buffer, pagesize=(width, height))
    if style == "tagvenue":
        _draw_tagvenue_overlay(pdf_canvas, row, width, height)
    elif style == "travis_perkins":
        _draw_travis_overlay(pdf_canvas, row, width, height)
    elif style == "united_protection":
        _draw_united_overlay(pdf_canvas, row, width, height)
    else:
        _draw_generic_overlay(pdf_canvas, row, width, height)
    pdf_canvas.save()
    buffer.seek(0)
    return PdfReader(buffer).pages[0]


def _detect_template_style(path: Path, text: str) -> str:
    haystack = f"{path.name}\n{text}".casefold()
    if "tagvenue" in haystack:
        return "tagvenue"
    if "travis" in haystack or "tp group" in haystack:
        return "travis_perkins"
    if "united protection" in haystack or "ups-" in haystack:
        return "united_protection"
    return "generic"


def _draw_tagvenue_overlay(canvas, row: dict[str, str], width: float, height: float) -> None:
    _white(canvas, 455, 762, 108, 24)
    _text(canvas, 458, 774, _value(row, "vat_invoice_number", "invoice_number"), size=9, bold=True)
    _white(canvas, 286, 633, 100, 18)
    _text(canvas, 286, 641, _value(row, "date_of_supply", "tax_point"), size=9)

    _white(canvas, 46, 585, 500, 92)
    _text(canvas, 50, 657, _value(row, "company_name", "supplier"), size=10, bold=True)
    _text(canvas, 50, 642, _value(row, "description", "specification"), size=9)
    _text(canvas, 50, 627, f"SKU: {_value(row, 'sku')}", size=9)
    _text(canvas, 390, 642, f"Qty: {_value(row, 'quantity')}", size=9)
    _text(canvas, 390, 627, f"Unit: {_value(row, 'unit_price')}", size=9)
    _text(canvas, 390, 612, f"VAT: {_value(row, 'vat_amount')}", size=9)

    _white(canvas, 395, 292, 170, 62)
    _text(canvas, 402, 334, f"Net: {_value(row, 'net_amount')}", size=10)
    _text(canvas, 402, 316, f"VAT: {_value(row, 'vat_amount')}", size=10)
    _text(canvas, 402, 298, f"Gross: {_value(row, 'gross_amount')}", size=11, bold=True)
    _white(canvas, 36, 230, 260, 60)
    _text(canvas, 40, 260, _value(row, "gross_amount"), size=26, bold=True)


def _draw_travis_overlay(canvas, row: dict[str, str], width: float, height: float) -> None:
    _white(canvas, 414, 705, 150, 64)
    _text(canvas, 418, 748, _value(row, "vat_invoice_number", "invoice_number"), size=11, bold=True)
    _text(canvas, 418, 732, _value(row, "payment_date", "invoice_date"), size=9)
    _text(canvas, 418, 716, _value(row, "date_of_supply", "tax_point"), size=9)

    _white(canvas, 54, 494, 488, 72)
    _text(canvas, 60, 546, _value(row, "sku"), size=8)
    _wrap_text(canvas, _value(row, "description", "specification"), 128, 546, 220, size=8)
    _text(canvas, 370, 546, _value(row, "quantity"), size=8)
    _text(canvas, 410, 546, _value(row, "unit_price"), size=8)
    _text(canvas, 488, 546, _value(row, "gross_amount", "net_amount"), size=8)

    _white(canvas, 378, 364, 170, 66)
    _text(canvas, 390, 410, f"Net: {_value(row, 'net_amount')}", size=9)
    _text(canvas, 390, 392, f"VAT: {_value(row, 'vat_amount')}", size=9)
    _text(canvas, 390, 374, f"Total: {_value(row, 'gross_amount')}", size=10, bold=True)


def _draw_united_overlay(canvas, row: dict[str, str], width: float, height: float) -> None:
    _white(canvas, 420, 720, 140, 38)
    _text(canvas, 424, 742, _value(row, "vat_invoice_number", "invoice_number"), size=11, bold=True)
    _text(canvas, 424, 726, _value(row, "payment_date", "invoice_date"), size=9)

    _white(canvas, 46, 610, 500, 95)
    _text(canvas, 52, 684, _value(row, "company_name", "supplier"), size=10, bold=True)
    _wrap_text(canvas, _value(row, "description", "specification"), 52, 666, 300, size=9)
    _text(canvas, 386, 666, f"Date: {_value(row, 'date_of_supply')}", size=9)
    _text(canvas, 386, 648, f"Qty: {_value(row, 'quantity')}", size=9)
    _text(canvas, 386, 630, f"Total: {_value(row, 'gross_amount')}", size=10, bold=True)

    _white(canvas, 380, 450, 165, 66)
    _text(canvas, 392, 496, f"Net: {_value(row, 'net_amount')}", size=9)
    _text(canvas, 392, 478, f"VAT: {_value(row, 'vat_amount')}", size=9)
    _text(canvas, 392, 460, f"Gross: {_value(row, 'gross_amount')}", size=10, bold=True)


def _draw_generic_overlay(canvas, row: dict[str, str], width: float, height: float) -> None:
    _white(canvas, 36, height - 250, width - 72, 180)
    _text(canvas, 48, height - 95, "Invoice data", size=16, bold=True)
    y = height - 120
    for key in (
        "company_name",
        "vat_invoice_number",
        "payment_date",
        "date_of_supply",
        "supplier",
        "sku",
        "description",
        "quantity",
        "unit_price",
        "net_amount",
        "vat_amount",
        "gross_amount",
    ):
        value = row.get(key, "")
        if not value:
            continue
        _text(canvas, 48, y, f"{key}: {value}", size=9)
        y -= 14


def _white(canvas, x: float, y: float, width: float, height: float) -> None:
    canvas.setFillColorRGB(1, 1, 1)
    canvas.setStrokeColorRGB(1, 1, 1)
    canvas.rect(x, y, width, height, fill=1, stroke=0)


def _text(canvas, x: float, y: float, text: str, *, size: int = 10, bold: bool = False) -> None:
    text = str(text or "")
    font_name = _font_name(text, bold=bold)
    canvas.setFillColorRGB(0.08, 0.12, 0.16)
    canvas.setFont(font_name, size)
    canvas.drawString(x, y, text)


def _wrap_text(canvas, text: str, x: float, y: float, max_width: float, *, size: int = 10) -> None:
    try:
        from reportlab.pdfbase.pdfmetrics import stringWidth
    except ImportError:
        _text(canvas, x, y, text, size=size)
        return
    font_name = _font_name(text)
    words = str(text or "").split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if current and stringWidth(candidate, font_name, size) > max_width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    for index, line in enumerate(lines[:4]):
        _text(canvas, x, y - index * (size + 3), line, size=size)


def _font_name(text: str, *, bold: bool = False) -> str:
    if any(ord(char) > 255 for char in str(text or "")):
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont

            if "STSong-Light" not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            return "STSong-Light"
        except Exception:
            return "Helvetica-Bold" if bold else "Helvetica"
    return "Helvetica-Bold" if bold else "Helvetica"


def _value(row: dict[str, str], *names: str) -> str:
    for name in names:
        normalized = normalize_invoice_header(name)
        value = row.get(normalized)
        if value:
            return value
    return ""
